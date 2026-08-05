from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


documents = Path(__file__).parents[1] / "public" / "documents"


# Форма № 10: значение категории годности должно находиться в нижней,
# обведённой ячейке пункта 5, а не над ней.
form10_path = documents / "form10-template.docx"
form10 = Document(form10_path)
section_two = form10.tables[18]
source_cell = section_two.cell(9, 1)
target_cell = section_two.cell(10, 1)
source_cell.text = ""
target_cell.text = "{{FITNESS_CATEGORY}}"
target_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for paragraph in target_cell.paragraphs:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
form10.save(form10_path)


# Справка Ф-2: дополняем верхний правый адресный блок постоянным текстом,
# сохраняя шрифт и оформление существующей строки «военного».
f2_path = documents / "f2-template.docx"
f2 = Document(f2_path)
address_cell = f2.tables[0].cell(0, 3)
updated = False
for paragraph in address_cell.paragraphs:
    for run in paragraph.runs:
        if "военного" in run.text:
            run.add_break()
            run.add_text("комиссариата г. Москва")
            updated = True
            break
    if updated:
        break
if not updated:
    raise RuntimeError("Строка «военного» в адресном блоке Ф-2 не найдена")
f2.save(f2_path)
